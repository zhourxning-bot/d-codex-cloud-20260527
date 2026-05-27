from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from pathlib import Path


OUT_DIR = Path("D:/codex/outputs")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "shopify独立站市场选择与避坑报告.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_table(table, widths, header=True):
    table.style = "Table Grid"
    set_table_width(table, widths)
    if header:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
        for cell in table.rows[0].cells:
            set_cell_shading(cell, "E8EEF5")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = "Microsoft YaHei"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    r.font.size = Pt(9.5)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                for r in p.runs:
                    r.font.name = "Microsoft YaHei"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    if r.font.size is None:
                        r.font.size = Pt(9.5)


def set_doc_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Shopify 独立站市场选择与避坑报告")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run("面向从 0 开始做跨境独立站的实操决策参考 | 2026-05-27")
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor.from_string("555555")


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [9360], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F3A5F")
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(2)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_market_table(doc):
    data = [
        ("1", "英国", "第一站优先", "英语市场、支付成熟、独立站接受度高；适合先验证产品、页面、投流和履约闭环。"),
        ("2", "澳大利亚", "第二复制市场", "客单价较好，竞争弱于美国；适合家居、户外、宠物、美妆个护、礼品。"),
        ("3", "美国", "模型跑通后放大", "赚钱人数最多、天花板最高；但广告贵、竞争强、税费和小包政策变化更明显。"),
        ("4", "德国", "品质货可做", "购买力强，但德语、VAT、EPR、产品合规和本地信任门槛更高。"),
        ("5", "加拿大", "北美副市场", "可复用美国素材和供应链，但单独体量小，不建议作为唯一主市场。"),
        ("6", "法国/意大利/西班牙", "本地化后再做", "有购买力，但语言、本地支付、退货和消费者信任要求更高。"),
        ("7", "沙特/UAE", "不建议新手首站", "有人靠 COD 和爆品赚钱，但拒收、回款、客服、清关和代理风险高。"),
        ("8", "墨西哥/巴西", "进阶市场", "增长快但支付、清关、税务、分期和本地物流复杂。"),
        ("9", "日本/韩国", "谨慎", "消费力强但信任门槛高，不适合粗糙页面和低质直发。"),
        ("10", "东南亚", "不优先", "人多但客单价低，平台和 TikTok/Shopee 更强，独立站起步并不轻松。"),
    ]
    table = doc.add_table(rows=1, cols=4)
    headers = ["排名", "市场", "建议", "原因"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in data:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    style_table(table, [650, 1250, 1700, 5760])


def add_loophole_table(doc):
    data = [
        ("支付风控", "多收款账户、多个主体轮换、突然放量后再换站", "账户关联、资金冻结 30-180 天、信用卡通道关闭", "从小额真实订单开始，保持同一主体、清晰售后、可追踪物流和低争议率。"),
        ("广告审核", "夸大功效、前后对比、假权威背书、落地页与广告不一致", "广告拒登、账户限制、像素数据报废", "素材和页面承诺一致，不做无法证明的健康、减肥、收益、效果承诺。"),
        ("素材版权", "搬运达人视频、Amazon/Temu/品牌图、竞品详情页", "DMCA 投诉、广告素材下架、品牌律师函", "优先自产图文视频，达人授权留存证据，供应商素材也要确认授权范围。"),
        ("评价与背书", "导入假五星评论、伪造买家图、虚构媒体报道", "消费者保护风险、广告平台风控、退款投诉", "用真实订单评价；没有证据的数字、奖项、媒体 logo 不要放。"),
        ("物流承诺", "写本地仓或 3-7 天，实际中国直发 10-25 天", "拒付率上升、PayPal/Stripe 风控、差评", "页面写真实处理时间和运输时间；能 DDP 就优先含税交付。"),
        ("产品合规", "补剂、美白、医疗器械、儿童用品、电子带电品未按当地规则处理", "广告封禁、清关扣货、召回或处罚", "新手先避开高监管品类，确定认证、标签、说明书、材质和安全标准后再卖。"),
        ("仿牌擦边", "蹭大牌词、近似 logo、同款外观、广告暗示替代品", "支付冻结、海关扣货、侵权索赔", "不要把侵权作为利润来源；产品、标题、素材、域名都避免借品牌势能。"),
        ("税与海关", "低申报、拆包、换站压低流水、伪装发货地", "追税、罚款、扣货、客户补税投诉", "把 VAT/GST/sales tax/关税当成本项；超过门槛就找税务代理处理。"),
        ("订阅与加购", "默认勾选保险/VIP、取消困难、试用后自动扣费不明显", "拒付、监管处罚、支付风控", "所有续费、加购、保险、服务费必须明确提示并可取消。"),
        ("数据营销", "买邮箱名单、短信轰炸、未同意再营销", "GDPR/CAN-SPAM/平台处罚、域名邮箱信誉受损", "只做同意基础上的邮件和短信营销，保留退订入口。"),
        ("客服售后", "模板拖延、失联、只退款不补发或只补发不退款", "争议集中爆发，支付账户被认为高风险", "设置 24-48 小时响应标准；物流异常、破损、错发要有固定处理规则。"),
        ("供应链", "爆单后换低配货、断货仍继续卖、样品和大货不一致", "货不对板、差评、退款潮", "先小批量验证质量；供应商、质检、备货、替代方案都要提前定。"),
        ("财务认知", "只看 ROAS，不算退款、拒付、税费、汇损、广告学习期", "以为赚钱，实际现金流亏损", "按单笔贡献利润算账，把冻结资金和退货损耗也算进去。"),
        ("代运营/培训", "承诺托管赚钱、展示不可验证流水、卖站卖账户", "被骗服务费、接手高风险资产", "只看可验证后台、合同、账户归属、历史争议率和真实物流记录。"),
    ]
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["风险口", "常见漏洞/灰色现象", "可能后果", "避坑做法"]):
        table.rows[0].cells[i].text = h
    for row in data:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    style_table(table, [1300, 2700, 2400, 2960])


def add_issue_table(doc):
    data = [
        ("Shopify/支付 reserve", "新店突然大量收款、退款/拒付高、履约慢、品类敏感", "先小额真实测试，物流追踪完整，退款政策清楚。"),
        ("PayPal 冻结", "新账户暴增、争议多、物流信息缺失", "每单及时上传追踪号，客服先处理再升级争议。"),
        ("Stripe/信用卡通道停用", "受限产品、误导页面、拒付率上升", "避开高风险品类，页面承诺真实，控制争议率。"),
        ("Meta/TikTok 广告不过审", "夸大效果、前后对比、擦边素材、落地页不一致", "先做合规素材库，广告词和详情页统一。"),
        ("Google Merchant 封禁", "商家信息不完整、价格库存不一致、退货政策不清", "补齐公司信息、联系方式、政策页和产品 feed。"),
        ("退货成本失控", "跨境退回贵、客户不愿承担、产品客单低", "低客单产品设置补发/部分退款策略，高客单建立本地退货点。"),
        ("毛利被吃掉", "税费、运费、支付手续费、退款、汇损没算", "先做完整单笔利润表，再决定能不能投广告。"),
        ("广告数据误判", "样本太小、像素没数据、只看点击不看结账", "至少看加购、发起结账、支付成功、退款和拒付。"),
        ("供应商质量波动", "样品好、大货差、批次不稳定", "首批小量，留样，要求发货前质检图/视频。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["问题", "常见触发", "预防动作"]):
        table.rows[0].cells[i].text = h
    for row in data:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    style_table(table, [1900, 3800, 3660])


def add_sources(doc):
    doc.add_heading("资料来源", level=1)
    sources = [
        "Shopify Help Center: Shopify Payments reserves, chargebacks, acceptable business practices.",
        "PayPal Help: account reserves and seller risk controls.",
        "Google Ads Policy: misrepresentation and unacceptable business practices.",
        "TikTok Ads Policy: misleading and false content.",
        "FTC: final rule banning fake reviews and testimonials; negative option/click-to-cancel guidance.",
        "EU OSS/IOSS, UK HMRC VAT guidance, Australian Border Force GST on low value goods.",
        "DHL cross-border ecommerce trends and public market data used as directional market context.",
        "说明：本报告是商业风险与入场策略参考，不构成法律、税务或会计意见；违规/灰色现象仅用于识别和避坑。",
    ]
    add_bullets(doc, sources)


def build():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc)

    add_callout(
        doc,
        "最终建议",
        "从 0 开始做 Shopify 独立站，优先英国单国站，小预算跑 2-4 周验证产品、页面、支付、物流和售后；跑通后复制澳大利亚，再考虑美国。不要把不交税、低申报、多账户轮换、假发货地等灰色漏洞当成商业模型。",
    )

    doc.add_heading("一、市场选择结论", level=1)
    doc.add_paragraph(
        "本报告按“实际卖家反复选择后仍愿意继续做、有人长期赚到钱、支付和广告可持续、履约风险可控”的角度排序，而不是只按国家电商规模排序。"
    )
    add_market_table(doc)

    doc.add_heading("二、建议起步路线", level=1)
    add_numbered(
        doc,
        [
            "第一站只做英国，不做全球站；货币、物流、政策页和客服语言都围绕英国。",
            "第 1-3 天完成 Shopify、域名、品牌邮箱、支付、政策页、产品页、物流说明。",
            "第 4-14 天做低强度养站：补内容、测下单、上传真实素材、引入少量真实访问。",
            "第 15-30 天小预算测试素材和产品角度，先看加购、结账、支付成功、退款咨询和物流反馈。",
            "30 天后若支付稳定、物流稳定、广告可控，再逐步加预算；同一模型再复制澳大利亚。",
            "美国放在模型跑通后作为放大市场；德国适合品质货，但必须补齐语言和合规能力。",
        ],
    )

    doc.add_heading("三、建站后马上用还是放一段时间", level=1)
    doc.add_paragraph(
        "建完空放意义不大。真正有价值的是域名、像素、广告账户、支付账户、内容、订单和用户行为的正常积累。最佳做法是建好后 7-14 天内开始小流量测试，30 天后再考虑正式放量。"
    )
    add_bullets(
        doc,
        [
            "不要一建站就大额投放：新站、新支付、新像素同时放量，容易触发广告和支付风控。",
            "不要空放几个月：没有访问、加购、订单和内容更新，站龄本身价值有限。",
            "养站的核心是正常经营痕迹：真实页面、真实素材、真实测试订单、清楚政策和稳定客服。",
        ],
    )

    doc.add_heading("四、税务与多站点现实", level=1)
    doc.add_paragraph(
        "独立站不是天然不用交税。一般会涉及卖家所在地所得税/企业税、客户所在地消费税，以及进口环节关税和 VAT/GST。早期确实有卖家不规范，但那是监管滞后，不是可持续优势。"
    )
    add_bullets(
        doc,
        [
            "多站点可以是正常经营：不同品牌、不同国家、不同品类、不同团队和真实业务线。",
            "多站点不能合法避税：如果目的是拆流水、换域名绕门槛、伪装发货地或规避申报，就是高风险行为。",
            "新手不要按“裸价利润”算账，必须把 VAT/GST/sales tax、关税、清关费、退款、拒付、汇损和支付费放进模型。",
        ],
    )

    doc.add_heading("五、行业里存在的漏洞与避坑清单", level=1)
    doc.add_paragraph(
        "下表列的是独立站圈里现实存在的空子、灰色操作或误区。目的不是照着做，而是让你在选服务商、选市场、选打法时能识别风险。"
    )
    add_loophole_table(doc)

    doc.add_heading("六、Shopify 新手最容易遇到的具体问题", level=1)
    add_issue_table(doc)

    doc.add_heading("七、品类选择建议", level=1)
    add_bullets(
        doc,
        [
            "优先：家居小工具、宠物用品、收纳、户外配件、礼品、非功效型美妆个护、轻定制产品。",
            "谨慎：带电产品、儿童用品、接触皮肤产品、食品接触类、汽车安全相关产品。",
            "新手先避开：补剂、医疗器械、减肥、美白、成人、仿牌、金融收益类、武器类、强功效承诺产品。",
            "选择标准：毛利至少能覆盖广告、运费、税费、退款损耗和支付手续费后仍有空间；供应链质量稳定，素材可自产。",
        ],
    )

    doc.add_heading("八、上线前检查清单", level=1)
    add_bullets(
        doc,
        [
            "域名、品牌邮箱、客服邮箱、退货地址或退货方案已准备好。",
            "退款政策、运输政策、隐私政策、服务条款、联系方式没有模板残留。",
            "产品页图片和视频有授权，文案没有夸大效果或虚假背书。",
            "物流时效写真实范围，不写无法兑现的本地仓和极速达。",
            "支付账户完成基本资料，测试订单和追踪号流程跑通。",
            "广告素材和落地页承诺一致，禁用前后对比、假医生、假媒体、假倒计时等高风险内容。",
            "单笔利润表包含产品成本、头程/尾程、税费、支付费、广告费、退款率、拒付率和汇损。",
            "客服模板覆盖：订单查询、物流延迟、破损、错发、退货、退款、取消订单。",
        ],
    )

    doc.add_heading("九、30 天执行计划", level=1)
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["阶段", "时间", "目标", "关键动作"]):
        table.rows[0].cells[i].text = h
    rows = [
        ("建站", "第 1-3 天", "基础可信", "Shopify、域名、支付、政策页、产品页、品牌邮箱、测试下单。"),
        ("养站", "第 4-14 天", "正常经营痕迹", "补 FAQ/指南、上传真实素材、小流量访问、检查移动端体验。"),
        ("测试", "第 15-30 天", "验证产品和页面", "小预算测 2-5 个素材角度，看加购、结账、支付和咨询。"),
        ("复盘", "第 30 天", "决定是否放量", "按利润表复盘，检查退款、物流、客服和支付风险。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    style_table(table, [1200, 1300, 2000, 4860])

    doc.add_heading("十、判断是否继续做的硬指标", level=1)
    add_bullets(
        doc,
        [
            "广告不是只看 ROAS，要看支付成功后的贡献利润。",
            "一周内如果咨询集中在物流和信任问题，优先修页面承诺和履约，不要盲目加预算。",
            "退款和拒付一旦上升，先暂停放量查原因；支付账户比广告数据更重要。",
            "如果一个产品必须靠虚假素材、夸大承诺或低申报才能赚钱，建议直接放弃。",
            "跑通英国后，再复制澳大利亚；复制时不要只翻货币，要重新检查物流、税费和退货成本。",
        ],
    )

    add_sources(doc)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("Shopify 独立站市场选择与避坑报告")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("777777")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
